"""CV upload + text extraction + structured parsing.

Supports PDF (pdfplumber/pypdf), DOCX (python-docx), TXT.
Falls back to plain-text heuristics if optional libs are missing.
"""
from __future__ import annotations

import io
import re
from typing import Any

# ==================== TEXT EXTRACTION ====================

_PDF_LIBS_TRIED: list[str] = []


def _extract_pdf(data: bytes) -> str:
    """Try pdfplumber first, then pypdf, then a best-effort stream parser."""
    if "pdfplumber" not in _PDF_LIBS_TRIED:
        try:
            import pdfplumber  # type: ignore

            with pdfplumber.open(io.BytesIO(data)) as pdf:
                txt = "\n".join((p.extract_text() or "") for p in pdf.pages)
            if txt.strip():
                return txt
        except Exception:
            _PDF_LIBS_TRIED.append("pdfplumber")
    if "pypdf" not in _PDF_LIBS_TRIED:
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(io.BytesIO(data))
            txt = "\n".join(p.extract_text() or "" for p in reader.pages)
            if txt.strip():
                return txt
        except Exception:
            _PDF_LIBS_TRIED.append("pypdf")
    # last resort: return bytes as latin-1
    return data.decode("latin-1", errors="ignore")


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document  # type: ignore

        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        # crude fallback: pull text between <w:t> tags
        return re.sub(r"<[^>]+>", " ", data.decode("utf-8", errors="ignore"))


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch on file extension."""
    fn = filename.lower()
    if fn.endswith(".pdf"):
        return _extract_pdf(data)
    if fn.endswith((".docx", ".doc")):
        return _extract_docx(data)
    if fn.endswith((".txt", ".md", ".text")):
        return data.decode("utf-8", errors="ignore")
    # unknown — try txt
    return data.decode("utf-8", errors="ignore")


# ==================== STRUCTURED PARSING ====================

_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(r"(\+?\d[\d\s().-]{7,}\d)")
_URL_RE = re.compile(r"https?://[^\s)>\]]+")

# common skill keyword hints (lowercase)
_SKILL_HINTS = {
    "python", "javascript", "typescript", "react", "node", "node.js", "go", "golang",
    "rust", "java", "kotlin", "swift", "ruby", "rails", "django", "flask", "fastapi",
    "postgresql", "postgres", "mysql", "mongodb", "redis", "elasticsearch", "kafka",
    "aws", "gcp", "azure", "kubernetes", "k8s", "docker", "terraform", "ansible",
    "pytorch", "tensorflow", "scikit-learn", "sklearn", "pandas", "numpy", "spark",
    "machine learning", "deep learning", "nlp", "computer vision", "llm", "mlops",
    "sql", "graphql", "rest", "grpc", "html", "css", "tailwind", "vue", "svelte",
    "ios", "android", "flutter", "react native", "next.js", "nuxt",
    "git", "ci/cd", "jenkins", "github actions", "circleci",
    "figma", "sketch", "adobe xd", "photoshop", "illustrator",
    "product management", "agile", "scrum", "kanban", "jira",
    "english", "urdu", "hindi", "arabic", "french", "german", "spanish",
}

_SECTION_HEADERS = {
    "experience": ["experience", "work experience", "professional experience", "employment", "employment history"],
    "education": ["education", "academic", "qualifications"],
    "skills": ["skills", "technical skills", "core competencies", "expertise"],
    "summary": ["summary", "profile", "about", "objective"],
}


def _split_sections(text: str) -> dict[str, str]:
    """Find sections by header keywords. Returns lowercase section name → text block."""
    lines = [ln.rstrip() for ln in text.splitlines()]
    out: dict[str, list[str]] = {}
    current: str | None = None
    for ln in lines:
        ll = ln.strip().lower().rstrip(":")
        matched_section: str | None = None
        if ll and len(ll) < 60:
            for sec, aliases in _SECTION_HEADERS.items():
                if ll in aliases or any(ll == a for a in aliases):
                    matched_section = sec
                    break
        if matched_section:
            current = matched_section
            out.setdefault(current, [])
        else:
            if current is None:
                current = "header"
                out.setdefault(current, [])
            out[current].append(ln)
    return {k: "\n".join(v).strip() for k, v in out.items()}


def _extract_email(text: str) -> str | None:
    m = _EMAIL_RE.search(text)
    return m.group(0) if m else None


def _extract_phone(text: str) -> str | None:
    m = _PHONE_RE.search(text)
    return m.group(0).strip() if m else None


def _extract_links(text: str) -> list[str]:
    return list(set(_URL_RE.findall(text)))[:10]


def _extract_name(text: str, email: str | None) -> str | None:
    """Heuristic: first non-empty line that looks like a name (2-4 words, title case)."""
    name_hint = email.split("@", 1)[0].replace(".", " ").replace("_", " ").replace("-", " ") if email else ""
    name_hint = " ".join(w for w in name_hint.split() if w and not w.isdigit()).title()
    for ln in text.splitlines()[:8]:
        s = ln.strip()
        if not s or "@" in s or "http" in s or len(s) > 60:
            continue
        parts = s.split()
        if 1 < len(parts) <= 4 and all(p[0:1].isupper() for p in parts if p):
            return s
    return name_hint or None


def _extract_skills(text: str) -> list[str]:
    """Return skills mentioned in the text (case-insensitive match against hint set)."""
    lower = " " + text.lower() + " "
    found: list[str] = []
    for skill in sorted(_SKILL_HINTS, key=len, reverse=True):
        # word-boundary-ish match
        if f" {skill} " in lower or f" {skill}," in lower or f" {skill}." in lower or f"({skill})" in lower:
            found.append(skill)
    return found[:30]


def _extract_experience(blocks: dict[str, str]) -> list[dict[str, Any]]:
    """Parse 'experience' section into a list of roles. Best-effort."""
    raw = blocks.get("experience", "").strip()
    if not raw:
        return []
    items: list[dict[str, Any]] = []
    # split on blank lines or "20XX" / "Jan 20XX" lines as a rough separator
    chunks = re.split(r"\n\s*\n", raw)
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk or len(chunk) < 10:
            continue
        lines = [l.strip() for l in chunk.splitlines() if l.strip()]
        first_line = lines[0]
        # try to find a date range on first 2 lines
        date_re = re.compile(
            r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{1,2}/)?\s*\d{4})\s*[\u2013\-]\s*"
            r"((?:Present|Current|Now|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{1,2}/)?\s*\d{4}))",
            re.IGNORECASE,
        )
        m = date_re.search(" ".join(lines[:2]))
        dates = m.group(0) if m else None
        # split first line by ' — ', ' at ', ' - ', ' | ' to get title/company
        title, company = None, None
        for sep in [" — ", " – ", " at ", " - ", " | ", ", "]:
            if sep in first_line:
                parts = first_line.split(sep, 1)
                if 1 < len(parts[0]) < 80:
                    title, company = parts[0].strip(), parts[1].strip()
                    break
        items.append(
            {
                "title": title or first_line,
                "company": company or "",
                "dates": dates,
                "description": " ".join(lines[1:])[:600],
            }
        )
    return items[:8]


def _extract_education(blocks: dict[str, str]) -> list[dict[str, Any]]:
    raw = blocks.get("education", "").strip()
    if not raw:
        return []
    items: list[dict[str, Any]] = []
    chunks = re.split(r"\n\s*\n", raw)
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        first_line = chunk.splitlines()[0].strip()
        # try to find year
        year_match = re.search(r"\b(19|20)\d{2}\b", first_line)
        items.append(
            {
                "school": first_line,
                "year": year_match.group(0) if year_match else None,
                "detail": " ".join(chunk.splitlines()[1:])[:300],
            }
        )
    return items[:5]


def parse_cv(filename: str, data: bytes) -> dict[str, Any]:
    """Extract text from a CV file and parse structured fields.

    Returns a dict with keys: name, email, phone, links, summary, skills, experience, education, raw_chars.
    """
    text = extract_text(filename, data)
    text = re.sub(r"[ \t]+", " ", text)  # collapse runs of spaces
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    blocks = _split_sections(text)
    email = _extract_email(text)
    phone = _extract_phone(text)
    name = _extract_name(text, email)
    links = _extract_links(text)
    skills = _extract_skills(text)
    experience = _extract_experience(blocks)
    education = _extract_education(blocks)
    summary = blocks.get("summary", "").strip() or blocks.get("header", "").strip()[:400] or None

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "links": links,
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "education": education,
        "raw_chars": len(text),
    }
